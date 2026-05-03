"""
Document Agent
==============
Generates a properly formatted .docx file from the document plan produced
by the FormattingAgent.

Features:
  - Per-paragraph font size, bold, italic, alignment, indentation.
  - Colour-coded headings that mirror professional Word styles.
  - Transparent agent metadata footer (explainability requirement).
  - Bullet-list support via Word's built-in List Bullet style.
  - Code paragraphs rendered in Courier New.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class DocumentAgent:
    """Generates a formatted .docx from an ordered document plan."""

    # Colour palette
    _COLOUR_TITLE      = RGBColor(0x1F, 0x35, 0x64)   # dark navy
    _COLOUR_HEADING    = RGBColor(0x1F, 0x35, 0x64)   # dark navy
    _COLOUR_SUBHEADING = RGBColor(0x2E, 0x74, 0xB5)   # mid blue
    _COLOUR_META       = RGBColor(0x99, 0x99, 0x99)   # grey

    def __init__(self, logger):
        self.logger = logger

    # ── Public API ────────────────────────────────────────────────────────────

    def generate(self, document_plan: list, output_path: str, analysis_meta: dict = None) -> str:
        """
        Write the .docx file and return its path.

        Parameters
        ----------
        document_plan : list
            Output of FormattingAgent.plan_document()
        output_path : str
            Destination file path (created/overwritten).
        analysis_meta : dict, optional
            Summary dict from the analysis (used for transparency header).
        """
        self.logger.log("DOCUMENT", f"Generating document → {os.path.basename(output_path)}")
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc = Document()
        self._set_margins(doc)

        # ── Transparency header ───────────────────────────────────────────────
        if analysis_meta:
            score = analysis_meta.get("confidence_score", 0) * 100
            self._add_meta_header(doc, score)

        # ── Body ─────────────────────────────────────────────────────────────
        for item in document_plan:
            self._add_paragraph(doc, item["text"], item["formatting"])

        doc.save(output_path)
        self.logger.log("DOCUMENT", "Document saved successfully.")
        return output_path

    # ── Paragraph builder ─────────────────────────────────────────────────────

    def _add_paragraph(self, doc: Document, text: str, fmt: dict):
        para_type = fmt.get("type", "body")

        if fmt.get("is_list"):
            para = doc.add_paragraph(style="List Bullet")
            run  = para.add_run(text.lstrip("•-*◦ "))
        else:
            para = doc.add_paragraph()
            run  = para.add_run(text)

        # Text run formatting
        run.bold         = fmt.get("bold",    False)
        run.italic       = fmt.get("italic",  False)
        run.font.size    = Pt(fmt.get("font_size", 11))

        # Alignment
        para.alignment = fmt.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)

        # Indentation
        level = fmt.get("indent_level", 0)
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.5 * level)

        # Type-specific decoration
        if para_type == "title":
            run.font.color.rgb = self._COLOUR_TITLE
            self._set_spacing(para, before=0, after=18, line=1.2)

        elif para_type == "heading":
            run.font.color.rgb = self._COLOUR_HEADING
            self._set_spacing(para, before=12, after=6)
            self._add_bottom_border(para, colour="4472C4")

        elif para_type == "subheading":
            run.font.color.rgb = self._COLOUR_SUBHEADING
            self._set_spacing(para, before=8, after=4)

        elif para_type == "code":
            run.font.name = "Courier New"
            self._set_spacing(para, before=4, after=4)

        elif para_type == "quote":
            run.italic                         = True
            para.paragraph_format.left_indent = Inches(0.5)
            self._set_spacing(para, before=4, after=4)

        else:
            self._set_spacing(para, before=3, after=3, line=1.15)

    # ── Document-level helpers ────────────────────────────────────────────────

    def _set_margins(self, doc: Document):
        section               = doc.sections[0]
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def _add_meta_header(self, doc: Document, score: float):
        para      = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run            = para.add_run(
            f"[Agentic OCR  ·  AI confidence: {score:.0f}%  ·  "
            "Observe → Interpret → Decide → Act → Learn]"
        )
        run.font.size      = Pt(8)
        run.font.color.rgb = self._COLOUR_META
        run.italic         = True
        self._set_spacing(para, before=0, after=12)

    # ── XML / OOXML helpers ───────────────────────────────────────────────────

    @staticmethod
    def _set_spacing(para, before: int = 6, after: int = 6, line: float = None):
        pPr     = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(before * 20))
        spacing.set(qn("w:after"),  str(after  * 20))
        if line is not None:
            spacing.set(qn("w:line"),     str(int(line * 240)))
            spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    @staticmethod
    def _add_bottom_border(para, colour: str = "4472C4"):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), colour)
        pBdr.append(bot)
        pPr.append(pBdr)
